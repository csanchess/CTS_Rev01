"""
Data Ingestion Processors
Handle various file formats and data sources
"""
from typing import Dict, Any, Optional
from fastapi import UploadFile
import pandas as pd
import io
from datetime import datetime
import uuid

from core.supabase_client import get_supabase_client


class DataIngestionProcessor:
    """Process various data formats"""
    
    def __init__(self):
        self.supabase = get_supabase_client()
    
    async def process_file(self, file: UploadFile, source_type: str) -> Dict[str, Any]:
        """Process an uploaded file"""
        ingestion_id = str(uuid.uuid4())
        
        try:
            # Create ingestion record
            self.supabase.table('data_ingestions').insert({
                'id': ingestion_id,
                'source_type': source_type,
                'source_name': file.filename,
                'file_name': file.filename,
                'status': 'processing',
                'created_at': datetime.utcnow().isoformat()
            }).execute()
            
            # Read file content
            content = await file.read()
            
            # Process based on type
            if source_type == 'spreadsheet':
                result = await self._process_spreadsheet(content, file.filename)
            elif source_type == 'pdf':
                result = await self._process_pdf(content)
            elif source_type == 'doc':
                result = await self._process_document(content)
            else:
                result = await self._process_text(content)
            
            # Update ingestion record
            self.supabase.table('data_ingestions').update({
                'status': 'completed',
                'records_processed': result.get('records_processed', 0),
                'metadata': result.get('metadata', {}),
                'completed_at': datetime.utcnow().isoformat()
            }).eq('id', ingestion_id).execute()
            
            return {
                'ingestion_id': ingestion_id,
                'records_processed': result.get('records_processed', 0),
                'metadata': result.get('metadata', {})
            }
            
        except Exception as e:
            # Update with error
            self.supabase.table('data_ingestions').update({
                'status': 'failed',
                'error_log': {'error': str(e)},
                'completed_at': datetime.utcnow().isoformat()
            }).eq('id', ingestion_id).execute()
            
            raise
    
    async def _process_spreadsheet(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process spreadsheet (Excel, CSV)"""
        try:
            # Determine file type
            if filename.endswith('.csv'):
                df = pd.read_csv(io.BytesIO(content))
            else:
                df = pd.read_excel(io.BytesIO(content))
            
            # Convert to records
            records = df.to_dict('records')
            
            # Here you would map columns to database schema
            # This is a simplified example
            
            return {
                'records_processed': len(records),
                'metadata': {
                    'columns': list(df.columns),
                    'row_count': len(records)
                }
            }
        except Exception as e:
            raise Exception(f"Error processing spreadsheet: {str(e)}")
    
    async def _process_pdf(self, content: bytes) -> Dict[str, Any]:
        """Process PDF file"""
        try:
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            
            text_content = ""
            for page in pdf_reader.pages:
                text_content += page.extract_text()
            
            # Process text content (extract entities, etc.)
            
            return {
                'records_processed': 1,
                'metadata': {
                    'pages': len(pdf_reader.pages),
                    'text_length': len(text_content)
                }
            }
        except Exception as e:
            raise Exception(f"Error processing PDF: {str(e)}")
    
    async def _process_document(self, content: bytes) -> Dict[str, Any]:
        """Process Word document"""
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            
            text_content = "\n".join([para.text for para in doc.paragraphs])
            
            return {
                'records_processed': 1,
                'metadata': {
                    'paragraphs': len(doc.paragraphs),
                    'text_length': len(text_content)
                }
            }
        except Exception as e:
            raise Exception(f"Error processing document: {str(e)}")
    
    async def _process_text(self, content: bytes) -> Dict[str, Any]:
        """Process plain text/log file"""
        try:
            text = content.decode('utf-8', errors='ignore')
            lines = text.split('\n')
            
            return {
                'records_processed': len(lines),
                'metadata': {
                    'lines': len(lines),
                    'text_length': len(text)
                }
            }
        except Exception as e:
            raise Exception(f"Error processing text: {str(e)}")
